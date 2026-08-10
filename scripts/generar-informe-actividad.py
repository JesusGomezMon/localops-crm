"""
Generate the Word activity report (design → TDD → CI/CD → simulated deploy).
Does not modify README.md.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CAPTURAS = ROOT / "entregables" / "capturas"
EVIDENCIA = ROOT / "entregables" / "evidencia"
OUT = ROOT / "entregables" / "Informe_Actividad_DevSecOps_Kasterz.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image(doc: Document, name: str, caption: str, width_in: float = 6.2) -> None:
    path = CAPTURAS / name
    if not path.exists():
        path = EVIDENCIA / name
    if not path.exists():
        add_p(doc, f"[Captura pendiente: {name}]")
        return
    doc.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Informe de actividad: desarrollo seguro, pruebas, CI/CD y despliegue simulado"
    )
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(
        f"Proyecto Kasterz (localops-crm) · Next.js 16 · Auth.js · Prisma/SQLite · "
        f"GitHub Actions · Docker Compose · {date.today().isoformat()}"
    )
    s.font.size = Pt(10)

    add_p(
        doc,
        "Este documento evidencia el uso de herramientas de desarrollo de software, "
        "plataformas de CI/CD, simulación de pruebas, entorno de despliegue y la "
        "aplicación de mejores prácticas de seguridad y privacidad. No forma parte "
        "del README del repositorio; se entrega como artefacto independiente en "
        "entregables/.",
    )

    # ------------------------------------------------------------------ 1
    add_heading(doc, "1. Diseño y planificación (seguridad y privacidad)", 1)
    add_heading(doc, "1.1 Objetivo", 2)
    add_p(
        doc,
        "Definir requisitos de la aplicación de reservas Kasterz con énfasis en "
        "seguridad y privacidad: panel de personal, reserva pública anónima, dos "
        "sucursales (Huayacán / Puerto Cancún) y pagos vía Stripe Checkout.",
    )

    add_heading(doc, "1.2 Requisitos de seguridad y privacidad documentados", 2)
    add_p(
        doc,
        "Se elaboraron y versionaron en el repositorio (docs/security-requirements.md, "
        "docs/threat-model.md, docs/architecture.md) los siguientes requisitos clave:",
    )
    add_bullets(
        doc,
        [
            "A1–A6 Autenticación: usuario + contraseña (scrypt), rate limit de login, "
            "mensajes de error uniformes, rechazo en producción de contraseña por defecto.",
            "Z1–Z3 Autorización: middleware + requireStaff(); reglas de sucursal VIP "
            "aplicadas en UI y servidor (isAvailableAtBranch).",
            "I1–I6 Integridad: Zod .strict(), precios solo en servidor, Prisma "
            "parametrizado, horarios fail-closed.",
            "P1–P4 Pagos: sin datos de tarjeta en la app; origen de redirect desde la "
            "URL del request; STRIPE_SECRET_KEY solo en entorno.",
            "D1–D4 Datos personales: mínimo (nombre, correo, teléfono); sin PII en "
            "respuestas de error al cliente.",
        ],
    )
    add_p(
        doc,
        "Referencias: OWASP ASVS / Top 10; NIST SP 800-63B (autenticación digital); "
        "principio de minimización del RGPD/LFPDPPP aplicable a datos de contacto.",
        bold=False,
    )

    add_heading(doc, "1.3 Análisis de amenazas (STRIDE) y defensas", 2)
    add_p(
        doc,
        "Se aplicó STRIDE sobre los activos: PII de clientes, cifras del negocio, "
        "calendario y clave Stripe. Mitigaciones implementadas incluyen:",
    )
    add_bullets(
        doc,
        [
            "Spoofing: credenciales con hash; rate limit por username; AUTH_SECRET para JWT.",
            "Tampering: totalCents no aceptado del cliente; extras con kind:'extra'; "
            "disponibilidad recalculada en el servidor.",
            "Information disclosure: errores genéricos; slots ocupados omitidos del calendario.",
            "DoS: rate limit de reservas públicas (BOOKING_RATE_LIMIT) y de login.",
            "Elevation: panel protegido; producción no arranca con admin123 por defecto.",
        ],
    )
    add_p(
        doc,
        "Referencia: Microsoft STRIDE; OWASP Threat Modeling Cheat Sheet.",
    )

    add_heading(doc, "1.4 Arquitectura de autenticación y autorización", 2)
    add_p(
        doc,
        "Dos poblaciones: clientes sin cuenta (reserva anónima) y personal con "
        "sesión JWT (Auth.js Credentials). La autorización del panel exige sesión; "
        "la regla de negocio de sucursal VIP es independiente de la autenticación.",
    )
    add_image(doc, "03-login.png", "Figura 1. Pantalla de acceso del personal (/login).")
    add_image(doc, "04-dashboard.png", "Figura 2. Panel tras autenticación (/dashboard).")

    # ------------------------------------------------------------------ 2
    add_heading(doc, "2. Desarrollo y pruebas", 1)
    add_heading(doc, "2.1 Objetivo", 2)
    add_p(
        doc,
        "Implementar funcionalidades clave con prácticas de codificación segura y "
        "verificarlas con pruebas automatizadas (enfoque TDD / pruebas primero en "
        "límites de seguridad).",
    )

    add_heading(doc, "2.2 Funcionalidades y codificación segura", 2)
    add_bullets(
        doc,
        [
            "Reserva pública: /book con catálogo filtrado por sucursal.",
            "Upsell y compra de membresías con checkout Stripe (precio desde BD).",
            "Panel: agenda, clientes, servicios, solicitudes de membresía.",
            "Validación Zod strict en fronteras de confianza; sin $queryRaw.",
            "Secretos fuera del repositorio (.env gitignored).",
        ],
    )
    add_image(doc, "01-inicio.png", "Figura 3. Entrada pública de la aplicación.")
    add_image(doc, "02-reservas.png", "Figura 4. Flujo de reserva pública.")
    add_image(doc, "05-servicios.png", "Figura 5. Catálogo de servicios (panel).")
    add_image(doc, "06-clientes.png", "Figura 6. Directorio de clientes (panel).")
    add_image(doc, "07-agenda.png", "Figura 7. Agenda y walk-in.")

    add_heading(doc, "2.3 Pruebas funcionales y de seguridad", 2)
    add_p(
        doc,
        "Suite Vitest en tests/: booking público, aislamiento por sucursal, "
        "anti-tampering de precios/extras, rate limiting (429 + Retry-After), "
        "compra de membresía, auth y dashboard. Equivalen a pruebas de regresión "
        "funcional y a controles de seguridad automatizados (no sustituyen un pentest "
        "externo completo, pero cubren los vectores documentados en el threat model).",
    )
    add_image(
        doc,
        "08-pruebas.png",
        "Figura 8. Ejecución de la suite de pruebas (pnpm test).",
    )
    add_p(
        doc,
        "Referencias: OWASP Testing Guide; pirámide de pruebas (unitarias/integración); "
        "TDD (Beck) — primero el comportamiento esperado del límite (p. ej. VIP branch "
        "rechaza Básico), luego la implementación.",
    )

    # ------------------------------------------------------------------ 3
    add_heading(doc, "3. Integración y entrega continua (CI/CD)", 1)
    add_heading(doc, "3.1 Objetivo", 2)
    add_p(
        doc,
        "Automatizar verificación en cada push/PR y preparar el artefacto "
        "contenedorizado para despliegue reproducible.",
    )

    add_heading(doc, "3.2 Pipeline CI (GitHub Actions)", 2)
    add_p(
        doc,
        "Archivo .github/workflows/ci.yml — job verify en ubuntu-latest:",
    )
    add_bullets(
        doc,
        [
            "checkout → pnpm + Node 20 con cache",
            "pnpm install --frozen-lockfile",
            "lint → typecheck → test → audit (nivel high) → build",
            "Variables de CI desechables (SQLite scratch, AUTH_SECRET de prueba)",
        ],
    )
    add_image(
        doc,
        "09-cicd-workflow.png",
        "Figura 9. Definición del workflow de CI en el repositorio.",
    )
    add_image(
        doc,
        "10-github-actions.png",
        "Figura 10. Ejecución del pipeline en GitHub Actions (tras push).",
    )
    add_p(
        doc,
        "Referencias: GitHub Actions security hardening; Dependabot/pnpm audit; "
        "principio fail-fast en el pipeline.",
    )

    add_heading(doc, "3.3 Contenedores y orquestación local", 2)
    add_p(
        doc,
        "Dockerfile multi-stage (deps → builder → runner Node 20 alpine, usuario "
        "no-root) y docker-compose.yml con servicios migrate (prisma migrate deploy "
        "+ seed) y app (standalone Next.js), volumen para SQLite. Compose es el "
        "entorno de orquestación documentado; cuando Docker Desktop no está "
        "disponible en el host, se usa el script scripts/deploy-simulado.ps1 que "
        "reproduce el mismo flujo en proceso Node de producción.",
    )
    add_image(
        doc,
        "11-docker-compose.png",
        "Figura 11. docker-compose.yml — migrate + app + volumen.",
    )
    add_p(
        doc,
        "Referencias: Docker multi-stage builds; CIS Docker Benchmark (usuario no-root, "
        "secretos por env); Twelve-Factor App (config en entorno).",
    )

    add_heading(doc, "3.4 Monitoreo y ajuste", 2)
    add_bullets(
        doc,
        [
            "Health check HTTP del despliegue simulado (GET / → 200).",
            "Logs de checkout Stripe y rate limit en stdout del proceso.",
            "Ajuste de BOOKING_RATE_LIMIT / BOOKING_RATE_WINDOW_MS por entorno.",
            "Auditoría de dependencias en CI (pnpm audit --audit-level high).",
        ],
    )

    # ------------------------------------------------------------------ 4
    add_heading(doc, "4. Lanzamiento y mantenimiento", 1)
    add_heading(doc, "4.1 Despliegue en entorno simulado", 2)
    add_p(
        doc,
        "Se desplegó la aplicación en modo producción simulado en "
        "http://localhost:3001 mediante scripts/deploy-simulado.ps1: migrate + seed "
        "con ADMIN_PASSWORD distinto al valor por defecto (exigido por el guard de "
        "producción en auth.ts), build y next start. Equivale funcionalmente al "
        "stack de docker compose up --build.",
    )
    add_image(
        doc,
        "12-despliegue-simulado.png",
        "Figura 12. Aplicación respondiendo en el entorno de despliegue simulado (:3001).",
    )

    add_heading(doc, "4.2 Retroalimentación de usuario (simulada) y ajustes", 2)
    add_bullets(
        doc,
        [
            "Hallazgo UX: tablas de Servicios/Clientes ilegibles (colores claros sobre tema oscuro) → corregido con tokens del design system.",
            "Hallazgo UX: login sin retorno a inicio → flecha atrás a /.",
            "Hallazgo producto: membresías solo como lead («te contactamos») → botón Comprar y checkout si hay datos de contacto.",
        ],
    )

    add_heading(doc, "4.3 Plan de mantenimiento y versionado", 2)
    add_bullets(
        doc,
        [
            "Control de versiones Git + GitHub; ramas feature / PR hacia main con CI obligatorio.",
            "Migraciones Prisma versionadas en prisma/migrations/.",
            "Rotación de ADMIN_PASSWORD y AUTH_SECRET antes de cualquier entorno real.",
            "Pendiente de producción: webhook Stripe, cifrado en reposo, borrado de PII bajo solicitud, roles owner/staff en acciones sensibles.",
            "SemVer del paquete (package.json 0.1.0) y etiquetas git para releases.",
        ],
    )
    add_p(
        doc,
        "Referencias: Semantic Versioning; Keep a Changelog; OWASP Application Security "
        "Verification Standard para backlog de hardening post-lanzamiento.",
    )

    # ------------------------------------------------------------------ 5
    add_heading(doc, "5. Evidencias de herramientas utilizadas", 1)
    add_bullets(
        doc,
        [
            "IDE / agente: Cursor — desarrollo, refactor y documentación.",
            "Runtime: Node.js ≥ 20.9, pnpm 11.9.",
            "Framework: Next.js 16 (App Router), TypeScript, Tailwind CSS v4.",
            "Auth: Auth.js v5 (Credentials + JWT).",
            "Datos: Prisma + SQLite.",
            "Pruebas: Vitest + Testing Library.",
            "CI/CD: GitHub Actions (.github/workflows/ci.yml).",
            "Contenedores: Dockerfile + docker-compose.yml; fallback scripts/deploy-simulado.ps1.",
            "Pagos (sandbox): Stripe Checkout vía REST/fetch.",
        ],
    )

    add_heading(doc, "6. Conclusión", 1)
    add_p(
        doc,
        "La actividad cubrió el ciclo completo: requisitos y amenazas, implementación "
        "con controles de seguridad, pruebas automatizadas, pipeline CI, empaquetado "
        "contenedorizado y despliegue en entorno simulado, con un plan de "
        "mantenimiento explícito. Las capturas adjuntas demuestran la aplicación "
        "en ejecución y las herramientas de verificación.",
    )

    doc.save(str(OUT))
    print(f"Documento generado: {OUT}")


if __name__ == "__main__":
    main()
